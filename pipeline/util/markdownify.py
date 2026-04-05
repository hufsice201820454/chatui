"""
파일 → 마크다운 변환 유틸
지원: PDF, DOCX, XLSX, HTML/HTM, TXT, MD
"""
import io
import logging
from typing import Optional

logger = logging.getLogger("pipeline.util.markdownify")


def to_markdown(data: bytes, file_name: str) -> str:
    """파일 bytes를 마크다운 텍스트로 변환."""
    ext = _ext(file_name)
    try:
        if ext == "pdf":
            return _pdf_to_markdown(data)
        elif ext == "docx":
            return _docx_to_markdown(data)
        elif ext == "xlsx":
            return _xlsx_to_markdown(data)
        elif ext in ("html", "htm"):
            return _html_to_markdown(data)
        elif ext in ("txt", "md"):
            return data.decode("utf-8", errors="replace")
        else:
            logger.warning("Unsupported format: %s", ext)
            return data.decode("utf-8", errors="replace")
    except Exception as e:
        logger.error("markdownify failed for %s: %s", file_name, e)
        return ""


def _ext(file_name: str) -> str:
    return file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""


def _pdf_to_markdown(data: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is required for PDF parsing. pip install pdfplumber")

    lines = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                lines.append(f"<!-- page {page_num} -->")
                lines.append(text)

            # 표 추출
            for table in page.extract_tables() or []:
                if not table:
                    continue
                header = table[0]
                rows = table[1:]
                md_table = _rows_to_md_table(header, rows)
                if md_table:
                    lines.append(md_table)

    return "\n\n".join(lines)


def _docx_to_markdown(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required. pip install python-docx")

    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading 1"):
            parts.append(f"# {text}")
        elif style.startswith("Heading 2"):
            parts.append(f"## {text}")
        elif style.startswith("Heading 3"):
            parts.append(f"### {text}")
        else:
            parts.append(text)

    # 표
    for table in doc.tables:
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        rows = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows[1:]
        ]
        md_table = _rows_to_md_table(header, rows)
        if md_table:
            parts.append(md_table)

    return "\n\n".join(parts)


def _xlsx_to_markdown(data: bytes) -> str:
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl is required. pip install openpyxl")

    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## {sheet_name}")
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        header = [str(c) if c is not None else "" for c in rows[0]]
        body = [
            [str(c) if c is not None else "" for c in row]
            for row in rows[1:]
        ]
        md_table = _rows_to_md_table(header, body)
        if md_table:
            parts.append(md_table)
    return "\n\n".join(parts)


def _html_to_markdown(data: bytes) -> str:
    try:
        import markdownify as mdlib
        html = data.decode("utf-8", errors="replace")
        return mdlib.markdownify(html, heading_style="ATX")
    except ImportError:
        # fallback: BeautifulSoup 텍스트 추출
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(data, "html.parser")
            return soup.get_text(separator="\n")
        except ImportError:
            return data.decode("utf-8", errors="replace")


def _rows_to_md_table(header: list, rows: list) -> str:
    if not header:
        return ""
    h = [str(c) if c is not None else "" for c in header]
    lines = ["| " + " | ".join(h) + " |", "| " + " | ".join(["---"] * len(h)) + " |"]
    for row in rows:
        cells = [str(c) if c is not None else "" for c in row]
        # 열 수 맞추기
        while len(cells) < len(h):
            cells.append("")
        lines.append("| " + " | ".join(cells[: len(h)]) + " |")
    return "\n".join(lines)
